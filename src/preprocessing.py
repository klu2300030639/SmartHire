import re
import os
import pypdf
import docx

# Custom list of standard English stopwords to avoid relying on nltk downloads
STOPWORDS = set([
    'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've", "you'll", "you'd",
    'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 'hers',
    'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what', 'which',
    'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been',
    'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if',
    'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between',
    'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out',
    'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why',
    'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not',
    'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't", 'should',
    "should've", 'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't",
    'didn', "didn't", 'doesn', "doesn't", 'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't",
    'ma', 'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't",
    'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', "wouldn't"
])

SKILLS_DICTIONARY = {
    "Software Engineering": [
        "java", "python", "c++", "spring boot", "rest apis", "microservices", "sql", "git", 
        "docker", "kubernetes", "system design", "multithreading", "postgresql", "redis", 
        "maven", "linux", "hibernate", "aws"
    ],
    "Data Science": [
        "python", "r", "sql", "machine learning", "deep learning", "statistics", "pandas", 
        "numpy", "scikit-learn", "pytorch", "tensorflow", "nlp", "tableau", "data visualization",
        "spark", "hadoop", "data mining", "jupyter"
    ],
    "Frontend Development": [
        "html5", "css3", "javascript", "typescript", "react", "angular", "vue.js", "tailwind css", 
        "sass", "ui/ux design", "ui/ux", "bootstrap", "webpack", "responsive design", "redux", "figma",
        "jest", "graphql", "next.js"
    ],
    "DevOps": [
        "aws", "azure", "gcp", "linux", "bash", "docker", "kubernetes", "terraform", "ansible", 
        "jenkins", "ci/cd", "prometheus", "grafana", "git", "nginx", "shell scripting",
        "cloudformation", "terraform cloud"
    ],
    "Product Management": [
        "product strategy", "roadmap", "agile", "scrum", "jira", "user research", "wireframing", 
        "product lifecycle", "sql", "a/b testing", "market analysis", "cross-functional leadership",
        "confluence", "data-driven decisions", "kpi tracking", "customer feedback"
    ],
    "Human Resources": [
        "recruiting", "talent acquisition", "employee relations", "hr policies", "onboarding", 
        "ats", "performance management", "payroll", "compliance", "sourcing", 
        "interviewing", "conflict resolution", "fmla", "hris", "diversity & inclusion"
    ]
}

def clean_text(text):
    """Cleans raw text by removing punctuation, email/phone, and excess whitespace."""
    if not text:
        return ""
    text = text.lower()
    # Replace newlines and tabs with space
    text = re.sub(r'[\r\n\t]+', ' ', text)
    # Remove emails first
    text = re.sub(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', ' ', text)
    # Remove URLs
    text = re.sub(r'http\S+', ' ', text)
    # Keep characters, numbers, and + # -
    text = re.sub(r'[^a-z0-9+#\-\s]', ' ', text)
    # Tokenize and remove stopwords
    words = text.split()
    cleaned_words = [w for w in words if w not in STOPWORDS]
    return " ".join(cleaned_words)

def extract_email(text):
    """Extracts email address from text."""
    if not text:
        return None
    match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', text)
    return match.group(0) if match else None

def extract_phone(text):
    """Extracts standard phone formats."""
    if not text:
        return None
    patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', # +1 555-555-5555
        r'(?:\+91[-.\s]?)?\d{5}[-.\s]?\d{5}',                    # +91 98765 43210
        r'\d{10}'                                             # 9876543210
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return None

def extract_skills(text):
    """Extracts skills from text based on skills dictionary."""
    if not text:
        return []
    text_lower = text.lower()
    found_skills = set()
    
    for category, skills in SKILLS_DICTIONARY.items():
        for skill in skills:
            if '++' in skill or '#' in skill:
                pattern = rf'\b{re.escape(skill)}'
            else:
                pattern = rf'\b{re.escape(skill)}\b'
                
            if re.search(pattern, text_lower):
                # Format skill nicely
                formatted_skill = skill.upper() if len(skill) <= 3 or skill in ['c++', 'c#', 'ats', 'hris', 'kpi'] else skill.title()
                found_skills.add(formatted_skill)
                
    return sorted(list(found_skills))

def parse_pdf(file_path):
    """Parses text from a PDF file."""
    try:
        reader = pypdf.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Error parsing PDF {file_path}: {e}")
        return ""

def parse_docx(file_path):
    """Parses text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        print(f"Error parsing DOCX {file_path}: {e}")
        return ""

def parse_txt(file_path):
    """Parses text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"Error parsing TXT {file_path}: {e}")
        return ""

def parse_resume(file_path):
    """Parses resume by delegating based on file type."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext == '.docx':
        return parse_docx(file_path)
    else:
        return parse_txt(file_path)
